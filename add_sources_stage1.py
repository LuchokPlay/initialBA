from pathlib import Path

from docx import Document


SOURCE = Path(r"F:/СПБГУ/курсовая/Лукин Егор Курсовая3.docx")
TARGET = Path(r"F:/СПБГУ/курсовая/Лукин Егор Курсовая3_источники_этап1.docx")


ADDITIONS = {
    "Многие компании работают с использованием информационных систем": (
        " Это соответствует логике управления бизнес-процессами, где анализируется не "
        "только выполнение отдельных операций, но и вся цепочка событий, действий и "
        "решений, формирующая результат процесса (Dumas et al., 2018)."
    ),
    "Предполагаемая модель данных представляет собой набор из экземпляров": (
        " Такая структура близка к представлению журнала событий в process mining: "
        "отдельный экземпляр процесса соответствует case, а действия с временными "
        "метками образуют последовательность выполнения процесса (van der Aalst, 2016)."
    ),
    "При предварительном анализе данных нужно будет установить": (
        " В литературе по process mining похожие задачи рассматриваются через анализ "
        "журналов событий, обнаружение процесса, проверку отклонений от ожидаемого "
        "поведения и исследование временной перспективы процесса "
        "(van der Aalst, 2016; van der Aalst et al., 2012)."
    ),
}

NEW_REFERENCES = [
    (
        "Dumas M., La Rosa M., Mendling J., Reijers H. A. Fundamentals of Business "
        "Process Management. 2nd ed. Cham: Springer, 2018. DOI: "
        "10.1007/978-3-662-56509-4."
    ),
    (
        "van der Aalst W. Process Mining: Data Science in Action. 2nd ed. Berlin, "
        "Heidelberg: Springer, 2016. DOI: 10.1007/978-3-662-49851-4."
    ),
    (
        "van der Aalst W. et al. Process Mining Manifesto // Business Process "
        "Management Workshops. Lecture Notes in Business Information Processing. "
        "Vol. 99. Berlin, Heidelberg: Springer, 2012. P. 169-194. DOI: "
        "10.1007/978-3-642-28108-2_19."
    ),
]


def append_if_missing(paragraph, addition):
    if addition not in paragraph.text:
        paragraph.add_run(addition)


def main():
    doc = Document(str(SOURCE))

    for marker, addition in ADDITIONS.items():
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                append_if_missing(paragraph, addition)
                break
        else:
            raise RuntimeError(f"Paragraph not found: {marker}")

    reference_title_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Список литературы":
            reference_title_index = index
            break
    if reference_title_index is None:
        raise RuntimeError("Reference section not found")

    existing_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for reference in NEW_REFERENCES:
        if reference not in existing_text:
            paragraph = doc.add_paragraph(reference)
            paragraph.style = "List Paragraph"

    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
